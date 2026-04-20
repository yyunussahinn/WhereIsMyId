"""
smart_checker.py
────────────────────────────────────────────────────────────────
Akıllı Tarama modülü.

Görevler:
  1. Appium'a bağlan, ekran görüntüsü al, elementleri topla (rect dahil)
  2. Driver'ı kapat
  3. claude_filter ile annotation kutularını eşleştir
  4. Word + Excel raporu üret

app.py'deki SmartTab tarafından thread içinde çağrılır.
Her log satırı log_cb(text) callback'i ile GUI'ye iletilir.
"""

import time
import os
from datetime import datetime
from collections import Counter

from claude_filter import filter_elements_by_boxes


# ════════════════════════════════════════════════════════════════════════════
#  PLATFORM SABITLERI
# ════════════════════════════════════════════════════════════════════════════

# iOS element tipleri
IOS_ALWAYS = [
    "XCUIElementTypeTextField",
    "XCUIElementTypeSecureTextField",
    "XCUIElementTypeButton",
    "XCUIElementTypeCell",
]
IOS_CONDITIONAL = ["XCUIElementTypeOther"]

# Android element tipleri
AND_ALWAYS = [
    "android.widget.EditText",
    "android.widget.Button",
    "android.widget.ImageButton",
    "android.widget.CheckBox",
    "android.widget.RadioButton",
    "android.widget.Switch",
    "android.widget.Spinner",
]
AND_CONDITIONAL = [
    "android.view.View",
    "android.view.ViewGroup",
    "android.widget.FrameLayout",
    "android.widget.LinearLayout",
    "android.widget.RelativeLayout",
    "android.widget.ImageView",
]
AND_RESOURCE_ONLY = ["android.widget.TextView"]

# Status sabitleri
STATUS_UNIQUE    = "ID Var"
STATUS_DUPLICATE = "Duplicate"
STATUS_MISSING   = "ID Yok"
STATUS_UNDEFINED = "Undefined ID"
NEW_STATUS_WAITING = "ID Eklenecek (Waiting Dev)"

STATUS_PALETTE = {
    STATUS_MISSING:   {"hdr": "C00000", "row": "FFDAD6", "alt": "FCEBEB", "txt": "501313"},
    STATUS_UNDEFINED: {"hdr": "C55A11", "row": "FCE4D6", "alt": "FFF3EC", "txt": "412402"},
    STATUS_DUPLICATE: {"hdr": "7B3F00", "row": "FAEEDA", "alt": "FEF6E4", "txt": "3B1F00"},
    STATUS_UNIQUE:    {"hdr": "375623", "row": "E2EFDA", "alt": "EAF3DE", "txt": "173404"},
}
NEW_STATUS_COLOR = {
    "hdr": "843C0C", "row": "FDE9D9", "alt": "FEF3EC", "txt": "843C0C",
}


def get_new_status(status: str) -> str:
    return "" if status == STATUS_UNIQUE else NEW_STATUS_WAITING


# ════════════════════════════════════════════════════════════════════════════
#  APPIUM BAĞLANTI & ELEMENT TOPLAMA
# ════════════════════════════════════════════════════════════════════════════

def connect_and_capture(platform: str, profile: dict,
                         appium_server: str, screenshot_path: str,
                         log_cb=print) -> tuple[list, str]:
    """
    Appium'a bağlan, screenshot al, elementleri topla, driver'ı kapat.

    Dönüş: (all_elements, detected_page)
    Her element dict'inde "rect" anahtarı var:
      {"x": int, "y": int, "width": int, "height": int}
    """
    from appium import webdriver
    from appium.webdriver.common.appiumby import AppiumBy

    log_cb("🚀 Appium driver başlatılıyor...")

    if platform == "ios":
        from appium.options.ios import XCUITestOptions
        options = XCUITestOptions()
        options.platform_name    = "iOS"
        options.device_name      = profile["device_name"]
        options.platform_version = profile["platform_version"]
        options.automation_name  = "XCUITest"
        options.bundle_id        = profile["bundle_id"]
        options.no_reset         = profile.get("no_reset", True)
        options.udid             = profile["udid"]
    else:
        from appium.options.android import UiAutomator2Options
        options = UiAutomator2Options()
        options.platform_name    = "Android"
        options.device_name      = profile["device_name"]
        options.platform_version = profile["platform_version"]
        options.automation_name  = "UiAutomator2"
        options.app_package      = profile["app_package"]
        options.app_activity     = profile["app_activity"]
        options.no_reset         = profile.get("no_reset", True)

    driver = webdriver.Remote(appium_server, options=options)
    time.sleep(3)

    log_cb("📸 Ekran görüntüsü alınıyor...")
    os.makedirs(os.path.dirname(screenshot_path), exist_ok=True)
    driver.get_screenshot_as_file(screenshot_path)
    log_cb(f"   → {screenshot_path}")

    detected_page = _get_detected_page(driver, platform)
    log_cb(f"   → Tespit edilen sayfa: {detected_page or '(bulunamadı)'}")

    log_cb("🔍 Elementler taranıyor...")
    screen_size = driver.get_window_size()
    sw = screen_size["width"]
    sh = screen_size["height"]

    all_elements = _collect_elements(driver, platform, detected_page, sw, sh,
                                      AppiumBy, log_cb)

    driver.quit()
    log_cb("✅ Driver kapatıldı.")

    # Özet
    counts = Counter(e["status"] for e in all_elements)
    log_cb(f"{'='*40}")
    log_cb(f"✅ Unique    : {counts[STATUS_UNIQUE]}")
    log_cb(f"⚠️  Undefined : {counts[STATUS_UNDEFINED]}")
    log_cb(f"🔁 Duplicate : {counts[STATUS_DUPLICATE]}")
    log_cb(f"❌ Missing   : {counts[STATUS_MISSING]}")
    log_cb(f"{'='*40}")

    return all_elements, detected_page


def _get_detected_page(driver, platform: str) -> str:
    try:
        if platform == "ios":
            import xml.etree.ElementTree as ET
            root = ET.fromstring(driver.page_source)
            for tag in ["XCUIElementTypeNavigationBar", "XCUIElementTypeStaticText"]:
                el = root.find(f".//{tag}")
                if el is not None:
                    lbl = el.get("label") or el.get("name") or ""
                    if lbl:
                        return lbl
        else:
            activity = driver.current_activity or ""
            return activity.split(".")[-1] if activity else ""
    except Exception:
        pass
    return ""


def _collect_elements(driver, platform: str, detected_page: str,
                       sw: int, sh: int, AppiumBy, log_cb) -> list:
    from appium.webdriver.common.appiumby import AppiumBy as AB

    candidates   = []
    all_elements = []

    if platform == "ios":
        all_elements, candidates = _collect_ios(
            driver, detected_page, sw, sh, AB)
    else:
        all_elements, candidates = _collect_android(
            driver, detected_page, AB)

    # Duplicate kontrolü
    name_counts = Counter(r["acc_id"] for r in candidates)
    for row in candidates:
        row["status"] = (STATUS_UNIQUE
                         if name_counts[row["acc_id"]] == 1
                         else STATUS_DUPLICATE)
        all_elements.append(row)

    return all_elements


def _collect_ios(driver, detected_page, sw, sh, AppiumBy):
    from appium.webdriver.common.appiumby import AppiumBy as AB

    def get_name(el):  return el.get_attribute("name")  or ""
    def get_label(el): return el.get_attribute("label") or ""
    def get_value(el): return el.get_attribute("value") or ""
    def short_type(t): return t.replace("XCUIElementType", "")

    def is_visible(el):
        try:
            r = el.rect
            w, h = r.get("width", 0), r.get("height", 0)
            x = r.get("x", 0)
            return w > 0 and h > 0 and x < sw and (x + w) > 0
        except Exception:
            return False

    def find_by_acc(name):
        try:
            driver.find_element(AB.ACCESSIBILITY_ID, name)
            return True
        except Exception:
            return False

    def has_real_id(name, label):
        return (name and name != label
                and not name.startswith("__")
                and find_by_acc(name))

    def is_undefined(name):
        return "undefined" in name.lower() or name.startswith("__")

    def get_rect(el):
        try:
            r = el.rect
            return {"x": r["x"], "y": r["y"],
                    "width": r["width"], "height": r["height"]}
        except Exception:
            return None

    all_elems  = []
    candidates = []

    for etype in IOS_ALWAYS + IOS_CONDITIONAL:
        elems = driver.find_elements(AB.XPATH, f"//{etype}")
        for el in elems:
            # interaktif kontrolü
            if etype in IOS_CONDITIONAL:
                if el.get_attribute("accessible") != "true":
                    continue
            if not is_visible(el):
                continue

            name    = get_name(el)
            label   = get_label(el)
            value   = get_value(el)
            display = label or value or ""
            stype   = short_type(etype)
            rect    = get_rect(el)

            if has_real_id(name, label):
                entry = {"page": detected_page, "type": stype,
                         "label": display, "value": value,
                         "acc_id": name, "rect": rect}
                if is_undefined(name):
                    entry["status"] = STATUS_UNDEFINED
                    all_elems.append(entry)
                else:
                    candidates.append(entry)
            else:
                all_elems.append({"page": detected_page, "type": stype,
                                   "label": display, "value": value,
                                   "acc_id": "", "status": STATUS_MISSING,
                                   "rect": rect})

    return all_elems, candidates


def _collect_android(driver, detected_page, AppiumBy):
    from appium.webdriver.common.appiumby import AppiumBy as AB

    def short_type(t):  return t.split(".")[-1]
    def clean(v):
        s = (v or "").strip()
        return "" if s.lower() in ("null", "none") else s

    def get_rid(el):
        rid = clean(el.get_attribute("resource-id"))
        if not rid:
            return ""
        return rid.split("/")[-1] if "/" in rid else rid

    def get_label(el):
        return clean(el.get_attribute("content-desc")) or clean(el.get_attribute("text"))

    def get_value(el):
        return clean(el.get_attribute("text"))

    def is_undefined(rid):
        return "undefined" in rid.lower()

    def is_interactive(el, etype):
        if etype in AND_ALWAYS:
            return True
        if etype in AND_RESOURCE_ONLY:
            return bool(get_rid(el))
        if etype in AND_CONDITIONAL:
            return (el.get_attribute("clickable") == "true"
                    or bool(get_rid(el)))
        return False

    def get_rect(el):
        try:
            r = el.rect
            return {"x": r["x"], "y": r["y"],
                    "width": r["width"], "height": r["height"]}
        except Exception:
            return None

    all_elems  = []
    candidates = []
    ALL_TYPES  = AND_ALWAYS + AND_CONDITIONAL + AND_RESOURCE_ONLY

    for etype in ALL_TYPES:
        elems = driver.find_elements(AB.XPATH, f"//{etype}")
        for el in elems:
            try:
                if not is_interactive(el, etype):
                    continue
                rid   = get_rid(el)
                label = get_label(el)
                value = get_value(el)
                stype = short_type(etype)
                rect  = get_rect(el)
            except Exception:
                continue

            if not rid and not label and not value:
                continue

            if rid:
                entry = {"page": detected_page, "type": stype,
                         "label": label, "value": value,
                         "acc_id": rid, "rect": rect}
                if is_undefined(rid):
                    entry["status"] = STATUS_UNDEFINED
                    all_elems.append(entry)
                else:
                    candidates.append(entry)
            else:
                all_elems.append({"page": detected_page, "type": stype,
                                   "label": label, "value": value,
                                   "acc_id": "", "status": STATUS_MISSING,
                                   "rect": rect})

    return all_elems, candidates


# ════════════════════════════════════════════════════════════════════════════
#  RAPOR ÜRETIMI
# ════════════════════════════════════════════════════════════════════════════

def generate_reports(elements: list, page_name: str,
                     output_dir: str, platform: str,
                     screenshot_path: str, output_fmt: str,
                     log_cb=print):
    """Word ve/veya Excel raporu üret."""

    os.makedirs(output_dir, exist_ok=True)
    plat_label = "iOS" if platform == "ios" else "ANDROID"
    plat_suffix = "IOS" if platform == "ios" else "Android"
    id_col_name = "Accessibility ID" if platform == "ios" else "Resource ID"

    word_file  = os.path.join(output_dir, f"{page_name}_smart_{plat_suffix}.docx")
    excel_file = os.path.join(output_dir, f"Smart_Report_{plat_suffix}.xlsx")

    if output_fmt in ("word", "word+excel"):
        _gen_word(elements, page_name, plat_label, id_col_name,
                  word_file, screenshot_path, log_cb)

    if output_fmt in ("excel", "word+excel"):
        _gen_excel(elements, page_name, plat_label, id_col_name,
                   excel_file, screenshot_path, log_cb)


def _gen_word(elements, page_name, plat_label, id_col_name,
              word_file, screenshot_path, log_cb):
    from docx import Document
    from docx.shared import RGBColor, Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from PIL import Image as PILImage

    def add_shading(cell, hex_color):
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), hex_color)
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:val"), "clear")
        cell._tc.get_or_add_tcPr().append(shading)

    def hex_to_rgb(h):
        return RGBColor(*bytes.fromhex(h))

    COLS     = ["Element ID", "Page", "Type", "Label / Text",
                "Value", id_col_name, "Status", "New Status"]
    COL_KEYS = ["element_id", "page", "type", "label",
                "value", "acc_id", "status", "new_status"]
    WIDTHS   = [Inches(1.2), Inches(0.8), Inches(0.9), Inches(1.3),
                Inches(0.9), Inches(1.4), Inches(0.8), Inches(1.5)]

    if os.path.exists(word_file):
        os.remove(word_file)
    doc = Document()

    title = doc.add_heading(
        f"Smart Accessibility Report — {page_name}", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dp = doc.add_paragraph(
        f"Tarih: {datetime.now().strftime('%d.%m.%Y %H:%M')}  |  Platform: {plat_label}")
    dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    if elements:
        table = doc.add_table(rows=1, cols=len(COLS))
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, col_name in enumerate(COLS):
            hdr[i].text = col_name
            run = hdr[i].paragraphs[0].runs[0]
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            hc = NEW_STATUS_COLOR["hdr"] if col_name == "New Status" else "2C2C2A"
            add_shading(hdr[i], hc)
            hdr[i].width = WIDTHS[i]

        for idx, elem in enumerate(elements):
            elem_id    = f"{page_name}_smart_{idx + 1}"
            status     = elem.get("status", STATUS_MISSING)
            new_status = get_new_status(status)
            palette    = STATUS_PALETTE.get(status, STATUS_PALETTE[STATUS_MISSING])
            row_hex    = palette["row"] if idx % 2 == 0 else palette["alt"]
            ns_hex     = NEW_STATUS_COLOR["row"] if idx % 2 == 0 else NEW_STATUS_COLOR["alt"]

            row_cells = table.add_row().cells
            for i, key in enumerate(COL_KEYS):
                val = (elem_id if key == "element_id"
                       else new_status if key == "new_status"
                       else elem.get(key, "") or "")
                row_cells[i].text  = val
                row_cells[i].width = WIDTHS[i]
                add_shading(row_cells[i], ns_hex if key == "new_status" else row_hex)
                runs = row_cells[i].paragraphs[0].runs
                if runs:
                    if key == "status":
                        runs[0].bold = True
                        runs[0].font.color.rgb = hex_to_rgb(palette["txt"])
                    elif key == "new_status" and new_status:
                        runs[0].bold = True
                        runs[0].font.color.rgb = hex_to_rgb(NEW_STATUS_COLOR["txt"])

    doc.add_paragraph("")

    if os.path.exists(screenshot_path):
        doc.add_heading("📸 Ekran Görüntüsü", level=2)
        with PILImage.open(screenshot_path) as img:
            w_px, _ = img.size
        w_in = min(w_px / 96, 5.5)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(screenshot_path, width=Inches(w_in))
        cap = doc.add_paragraph(f"{page_name} sayfası ekran görüntüsü")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].font.size = Pt(9)
        cap.runs[0].italic    = True

    doc.save(word_file)
    log_cb(f"📄 Word kaydedildi: {word_file}")


def _gen_excel(elements, page_name, plat_label, id_col_name,
               excel_file, screenshot_path, log_cb):
    import openpyxl
    from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    from openpyxl.drawing.image import Image as XLImage
    from PIL import Image as PILImage

    THIN   = Side(style="thin")
    BORDER = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)
    CENTER = Alignment(horizontal="center", vertical="center", wrap_text=True)
    LEFT   = Alignment(horizontal="left",   vertical="center", wrap_text=True)

    COLS     = ["Element ID", "Page", "Type", "Label / Text",
                "Value", id_col_name, "Status", "New Status"]
    COL_KEYS = ["element_id", "page", "type", "label",
                "value", "acc_id", "status", "new_status"]
    WIDTHS   = [22, 16, 16, 26, 18, 32, 14, 28]

    DATA_COL = len(COLS)
    IMG_COL  = DATA_COL + 2
    IMG_LTR  = get_column_letter(IMG_COL)

    wb = (openpyxl.load_workbook(excel_file)
          if os.path.exists(excel_file) else openpyxl.Workbook())
    if not os.path.exists(excel_file) and "Sheet" in wb.sheetnames:
        del wb["Sheet"]
    if page_name in wb.sheetnames:
        del wb[page_name]
    ws = wb.create_sheet(title=page_name)

    # Başlık
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=DATA_COL)
    c = ws.cell(row=1, column=1,
                value=f"{page_name}  |  {datetime.now().strftime('%d.%m.%Y %H:%M')}  |  {plat_label}  |  Smart")
    c.font      = Font(bold=True, color="FFFFFF", size=13)
    c.fill      = PatternFill("solid", fgColor="1F3864")
    c.alignment = CENTER
    c.border    = BORDER
    ws.row_dimensions[1].height = 26

    # Kolon başlıkları
    for ci, col_name in enumerate(COLS, 1):
        c = ws.cell(row=2, column=ci, value=col_name)
        c.font = Font(bold=True, color="FFFFFF", size=10)
        hc = NEW_STATUS_COLOR["hdr"] if col_name == "New Status" else "2C2C2A"
        c.fill = PatternFill("solid", fgColor=hc)
        c.alignment = CENTER
        c.border    = BORDER
    ws.row_dimensions[2].height = 18
    ws.freeze_panes = "A3"

    # Veri
    for idx, elem in enumerate(elements):
        elem_id    = f"{page_name}_smart_{idx + 1}"
        status     = elem.get("status", STATUS_MISSING)
        new_status = get_new_status(status)
        row_num    = idx + 3
        palette    = STATUS_PALETTE.get(status, STATUS_PALETTE[STATUS_MISSING])
        rf = PatternFill("solid", fgColor=palette["row"] if idx % 2 == 0 else palette["alt"])
        nf = PatternFill("solid", fgColor=NEW_STATUS_COLOR["row"] if idx % 2 == 0 else NEW_STATUS_COLOR["alt"])

        for ci, key in enumerate(COL_KEYS, 1):
            val = (elem_id if key == "element_id"
                   else new_status if key == "new_status"
                   else elem.get(key, "") or "")
            c = ws.cell(row=row_num, column=ci, value=val)
            c.border = BORDER
            if key == "new_status":
                c.fill = nf
                c.font = Font(bold=bool(new_status), color=NEW_STATUS_COLOR["txt"], size=10)
                c.alignment = CENTER
            elif key == "status":
                c.fill = rf
                c.font = Font(bold=True, color=palette["txt"], size=10)
                c.alignment = CENTER
            elif key == "element_id":
                c.fill = rf
                c.font = Font(bold=True, size=10)
                c.alignment = CENTER
            else:
                c.fill = rf
                c.font = Font(size=10)
                c.alignment = LEFT
        ws.row_dimensions[row_num].height = 16

    for ci, w in enumerate(WIDTHS, 1):
        ws.column_dimensions[get_column_letter(ci)].width = w

    # Screenshot
    if os.path.exists(screenshot_path):
        with PILImage.open(screenshot_path) as img:
            ow, oh = img.size
        tw = 300
        th = int(oh * (tw / ow))
        tmp = screenshot_path.replace(".png", "_xl_tmp.png")
        with PILImage.open(screenshot_path) as img:
            img.resize((tw, th), PILImage.LANCZOS).save(tmp, "PNG")

        ws.column_dimensions[get_column_letter(DATA_COL + 1)].width = 2
        ws.merge_cells(start_row=1, start_column=IMG_COL,
                       end_row=2,   end_column=IMG_COL)
        hc = ws.cell(row=1, column=IMG_COL, value=f"📸 {page_name}")
        hc.font = Font(bold=True, color="FFFFFF", size=10)
        hc.fill = PatternFill("solid", fgColor="1F3864")
        hc.alignment = CENTER
        hc.border    = BORDER
        ws.column_dimensions[IMG_LTR].width = 42

        xi = XLImage(tmp)
        xi.width = tw; xi.height = th
        ws.add_image(xi, f"{IMG_LTR}3")

    wb.save(excel_file)
    log_cb(f"📊 Excel kaydedildi: {excel_file}  (sheet: {page_name})")

    if os.path.exists(screenshot_path):
        try:
            os.remove(tmp)
        except Exception:
            pass