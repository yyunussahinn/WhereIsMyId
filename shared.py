"""
shared.py — Where is My Id
────────────────────────────────────────────────────────────────────────────
Ortak sabitler, renk paleti, stil yardımcıları ve çıktı üreticileri.
element_checker_ios.py, element_checker_android.py ve build_summary.py
tarafından import edilir.
"""

import os
import sys
import tempfile
from datetime import datetime

import openpyxl
from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
from openpyxl.utils import get_column_letter

# ── Status sabitleri ──────────────────────────────────────────────────────────
STATUS_UNIQUE    = "ID Var"
STATUS_DUPLICATE = "Duplicate"
STATUS_MISSING   = "ID Yok"
STATUS_UNDEFINED = "Undefined ID"
ALL_STATUSES     = [STATUS_MISSING, STATUS_UNDEFINED, STATUS_DUPLICATE, STATUS_UNIQUE]

NS_WAITING = "ID Eklenecek (Waiting Dev)"

SECTION_TO_STATUS: dict[str, str] = {
    "missing":   STATUS_MISSING,
    "undefined": STATUS_UNDEFINED,
    "duplicate": STATUS_DUPLICATE,
    "unique":    STATUS_UNIQUE,
}

def get_new_status(status: str) -> str:
    return "" if status == STATUS_UNIQUE else NS_WAITING


# ── Renk paleti ───────────────────────────────────────────────────────────────
STATUS_PALETTE: dict[str, dict] = {
    STATUS_MISSING:   {"hdr": "C00000", "row": "FFDAD6", "alt": "FCEBEB", "txt": "501313"},
    STATUS_UNDEFINED: {"hdr": "C55A11", "row": "FCE4D6", "alt": "FFF3EC", "txt": "412402"},
    STATUS_DUPLICATE: {"hdr": "7B3F00", "row": "FAEEDA", "alt": "FEF6E4", "txt": "3B1F00"},
    STATUS_UNIQUE:    {"hdr": "375623", "row": "E2EFDA", "alt": "EAF3DE", "txt": "173404"},
}
NEW_STATUS_COLOR: dict[str, str]    = {"hdr": "843C0C", "row": "FDE9D9", "alt": "FEF3EC", "txt": "843C0C"}
AI_SUGGESTION_COLOR: dict[str, str] = {"hdr": "1F4E79", "row": "DEEAF1", "alt": "EBF3F9", "txt": "1F4E79"}

# ── openpyxl stil sabitleri ───────────────────────────────────────────────────
_THIN  = Side(style="thin")
BORDER = Border(left=_THIN, right=_THIN, top=_THIN, bottom=_THIN)
CENTER = Alignment(horizontal="center", vertical="center", wrap_text=True)
LEFT   = Alignment(horizontal="left",   vertical="center", wrap_text=True)

def fill(hex_c: str) -> PatternFill:
    return PatternFill("solid", fgColor=hex_c)

def hdr_font(size: int = 10) -> Font:
    return Font(bold=True, color="FFFFFF", size=size)


# ── Güvenli Excel kaydı ───────────────────────────────────────────────────────
def safe_save(wb: openpyxl.Workbook, filepath: str) -> None:
    """
    Workbook'u önce geçici dosyaya kaydeder, ardından atomik os.replace() ile
    hedefe taşır.

    Neden gerekli?
    • OneDrive / iCloud klasörlerine doğrudan wb.save(path) yapıldığında,
      sync daemon dosyayı kilitlemiş ya da yeni sürüm yüklüyorsa macOS
      çakışma kopyası oluşturur → "Elements_Report_Android-N100165.xlsx" gibi.
    • Geçici dosya → atomik rename yöntemi bu pencereyi en aza indirir.
    • Dosya hâlâ açıksa (Numbers / Excel) PermissionError fırlatır; kullanıcı
      bilgilendirilir.
    """
    dirpath = os.path.dirname(os.path.abspath(filepath))
    fd, tmp = tempfile.mkstemp(suffix=".xlsx", dir=dirpath)
    os.close(fd)
    try:
        wb.save(tmp)
        os.replace(tmp, filepath)   # POSIX'te atomik
    except PermissionError:
        try:
            os.remove(tmp)
        except OSError:
            pass
        raise PermissionError(
            f"❌ '{os.path.basename(filepath)}' dosyası başka bir uygulama tarafından "
            "açık. Lütfen Excel / Numbers'ı kapatıp tekrar çalıştırın."
        )
    except Exception:
        try:
            os.remove(tmp)
        except OSError:
            pass
        raise


# ── Kullanıcı onay soruları ───────────────────────────────────────────────────
def ask_overwrite(label: str) -> bool:
    """[e/h] sorusu; doğru cevap gelene kadar tekrar eder."""
    while True:
        sys.stdout.flush()
        ans = input(
            f"   ⚠️  {label} zaten mevcut. Üzerine yazmak istiyor musunuz? [e/h]: "
        ).strip().lower()
        if ans in ("e", "evet", "y", "yes"):  return True
        if ans in ("h", "hayır", "n", "no"): return False
        print("   Lütfen 'e' (evet) veya 'h' (hayır) girin.")


# ── AI Suggestion zenginleştirici ─────────────────────────────────────────────
def enrich_with_ai(all_elements: list, platform: str) -> list:
    """
    ai_suggestion.py mevcutsa elementleri zenginleştirir.
    Yoksa veya hata verirse 'ai_suggestion' anahtarını boş bırakır.
    """
    try:
        from ai_suggestion import enrich_elements
        return enrich_elements(all_elements, platform)
    except ImportError:
        print("⚠️  ai_suggestion.py bulunamadı. AI Suggestion sütunu boş kalacak.")
    except Exception as ex:
        print(f"⚠️  AI Suggestion hatası: {ex}. Sütun boş kalacak.")
    for e in all_elements:
        e.setdefault("ai_suggestion", "")
    return all_elements


# ── Excel çıktı üreticisi ─────────────────────────────────────────────────────
def generate_excel(
    all_elements:      list,
    page_name:         str,
    excel_file:        str,
    document_sections: list,
    platform:          str,        # "ios" | "android"
    screenshot_path:   str = "",
) -> None:
    from openpyxl.drawing.image import Image as XLImage
    from PIL import Image as PILImage

    acc_col  = "Accessibility ID" if platform == "ios" else "Resource ID"
    COLS     = ["Element ID", "Page", "Type", "Label / Text", "Value",
                acc_col, "Status", "New Status", "AI Suggestion"]
    COL_KEYS = ["element_id", "page", "type", "label", "value",
                "acc_id",     "status", "new_status", "ai_suggestion"]
    WIDTHS   = [22, 16, 16, 26, 18, 32, 14, 28, 45]

    DATA_COLS  = len(COLS)
    IMG_COL    = DATA_COLS + 2                       # boşluk sütunu atlayarak
    IMG_LTR    = get_column_letter(IMG_COL)
    PLAT_LABEL = "iOS" if platform == "ios" else "ANDROID"

    # Mevcut workbook'u yükle veya yeni oluştur
    wb = (openpyxl.load_workbook(excel_file)
          if os.path.exists(excel_file) else openpyxl.Workbook())
    # Boş varsayılan sheet'i temizle
    if "Sheet" in wb.sheetnames and len(wb.sheetnames) == 1:
        del wb["Sheet"]
    # Aynı isimli sheet varsa sil
    if page_name in wb.sheetnames:
        del wb[page_name]
    ws = wb.create_sheet(title=page_name)

    # Başlık satırı
    ts = datetime.now().strftime("%d.%m.%Y %H:%M")
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=DATA_COLS)
    c = ws.cell(row=1, column=1, value=f"{page_name}  |  {ts}  |  {PLAT_LABEL}")
    c.font = Font(bold=True, color="FFFFFF", size=13)
    c.fill = fill("1F3864"); c.alignment = CENTER; c.border = BORDER
    ws.row_dimensions[1].height = 26

    # Kolon başlıkları
    for ci, col_name in enumerate(COLS, 1):
        c = ws.cell(row=2, column=ci, value=col_name)
        c.font = hdr_font()
        if col_name == "AI Suggestion":   hdr_c = AI_SUGGESTION_COLOR["hdr"]
        elif col_name == "New Status":    hdr_c = NEW_STATUS_COLOR["hdr"]
        else:                             hdr_c = "2C2C2A"
        c.fill = fill(hdr_c); c.alignment = CENTER; c.border = BORDER
    ws.row_dimensions[2].height = 18
    ws.freeze_panes = "A3"

    # Veri satırları
    ordered = _build_ordered(all_elements, document_sections)
    for idx, elem in enumerate(ordered):
        elem_id    = f"{page_name}_element_{idx + 1}"
        status     = elem.get("status", STATUS_MISSING)
        new_status = get_new_status(status)
        row_num    = 3 + idx
        pal        = STATUS_PALETTE.get(status, STATUS_PALETTE[STATUS_MISSING])
        r_fill     = fill(pal["row"] if idx % 2 == 0 else pal["alt"])
        ns_fill    = fill(NEW_STATUS_COLOR["row"] if idx % 2 == 0 else NEW_STATUS_COLOR["alt"])
        ai_fill    = fill(AI_SUGGESTION_COLOR["row"] if idx % 2 == 0 else AI_SUGGESTION_COLOR["alt"])

        for ci, key in enumerate(COL_KEYS, 1):
            val = _get_val(elem, key, elem_id, new_status)
            c   = ws.cell(row=row_num, column=ci, value=val)
            c.border = BORDER
            _style_excel_cell(c, key, pal, r_fill, ns_fill, ai_fill, new_status)

        ws.row_dimensions[row_num].height = 60

    for ci, w in enumerate(WIDTHS, 1):
        ws.column_dimensions[get_column_letter(ci)].width = w

    # Ekran görüntüsü
    tmp_path = None
    if screenshot_path and os.path.exists(screenshot_path):
        try:
            with PILImage.open(screenshot_path) as img:
                orig_w, orig_h = img.size
            target_w = 300
            target_h = int(orig_h * target_w / orig_w)
            tmp_path = screenshot_path.replace(".png", "_xl_tmp.png")
            with PILImage.open(screenshot_path) as img:
                img.resize((target_w, target_h), PILImage.LANCZOS).save(tmp_path, "PNG")

            ws.column_dimensions[get_column_letter(DATA_COLS + 1)].width = 2
            ws.merge_cells(start_row=1, start_column=IMG_COL, end_row=2, end_column=IMG_COL)
            hc = ws.cell(row=1, column=IMG_COL, value=f"📸 {page_name}")
            hc.font = hdr_font(); hc.fill = fill("1F3864")
            hc.alignment = CENTER; hc.border = BORDER
            ws.column_dimensions[IMG_LTR].width = 42

            xl_img = XLImage(tmp_path)
            xl_img.width = target_w; xl_img.height = target_h
            ws.add_image(xl_img, f"{IMG_LTR}3")
        except Exception as e:
            print(f"⚠️  Ekran görüntüsü Excel'e eklenemedi: {e}")

    safe_save(wb, excel_file)
    print(f"📊 Excel kaydedildi: {excel_file}  (sheet: {page_name})")

    if tmp_path and os.path.exists(tmp_path):
        try:
            os.remove(tmp_path)
        except OSError:
            pass


# ── Word çıktı üreticisi ──────────────────────────────────────────────────────
def generate_word(
    all_elements:      list,
    page_name:         str,
    word_file:         str,
    document_sections: list,
    platform:          str,
    screenshot_path:   str = "",
) -> None:
    from docx import Document
    from docx.shared import RGBColor, Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from PIL import Image as PILImage

    acc_col    = "Accessibility ID" if platform == "ios" else "Resource ID"
    PLAT_LABEL = "iOS" if platform == "ios" else "ANDROID"
    COLS       = ["Element ID", "Page", "Type", "Label / Text", "Value",
                  acc_col, "Status", "New Status", "AI Suggestion"]
    COL_KEYS   = ["element_id", "page", "type", "label", "value",
                  "acc_id",     "status", "new_status", "ai_suggestion"]
    WIDTHS     = [Inches(1.0), Inches(0.7), Inches(0.8), Inches(1.1),
                  Inches(0.8), Inches(1.2), Inches(0.8), Inches(1.3), Inches(2.0)]

    def _shading(cell, hex_color: str) -> None:
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), hex_color)
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:val"), "clear")
        cell._tc.get_or_add_tcPr().append(shd)

    def _rgb(h: str) -> RGBColor:
        return RGBColor(*bytes.fromhex(h))

    if os.path.exists(word_file):
        os.remove(word_file)
    doc = Document()

    title      = doc.add_heading(f"Accessibility ID Report — {page_name}", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dp = doc.add_paragraph(
        f"Tarih: {datetime.now().strftime('%d.%m.%Y %H:%M')}  |  Platform: {PLAT_LABEL}")
    dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    ordered = _build_ordered(all_elements, document_sections)
    if ordered:
        table = doc.add_table(rows=1, cols=len(COLS))
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, col_name in enumerate(COLS):
            hdr[i].text = col_name
            run = hdr[i].paragraphs[0].runs[0]
            run.bold = True; run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            if col_name == "AI Suggestion":   hdr_c = AI_SUGGESTION_COLOR["hdr"]
            elif col_name == "New Status":    hdr_c = NEW_STATUS_COLOR["hdr"]
            else:                             hdr_c = "2C2C2A"
            _shading(hdr[i], hdr_c)
            hdr[i].width = WIDTHS[i]

        for idx, elem in enumerate(ordered):
            elem_id    = f"{page_name}_element_{idx + 1}"
            status     = elem.get("status", STATUS_MISSING)
            new_status = get_new_status(status)
            pal        = STATUS_PALETTE.get(status, STATUS_PALETTE[STATUS_MISSING])
            row_hex    = pal["row"] if idx % 2 == 0 else pal["alt"]
            ns_hex     = NEW_STATUS_COLOR["row"] if idx % 2 == 0 else NEW_STATUS_COLOR["alt"]
            ai_hex     = AI_SUGGESTION_COLOR["row"] if idx % 2 == 0 else AI_SUGGESTION_COLOR["alt"]

            row_cells = table.add_row().cells
            for i, key in enumerate(COL_KEYS):
                val = _get_val(elem, key, elem_id, new_status)
                row_cells[i].text  = val
                row_cells[i].width = WIDTHS[i]

                if key == "ai_suggestion":  _shading(row_cells[i], ai_hex)
                elif key == "new_status":   _shading(row_cells[i], ns_hex)
                else:                       _shading(row_cells[i], row_hex)

                runs = row_cells[i].paragraphs[0].runs
                if runs:
                    if key == "status":
                        runs[0].bold = True
                        runs[0].font.color.rgb = _rgb(pal["txt"])
                    elif key == "new_status" and new_status:
                        runs[0].bold = True
                        runs[0].font.color.rgb = _rgb(NEW_STATUS_COLOR["txt"])
                    elif key == "ai_suggestion" and val:
                        runs[0].font.size      = Pt(7)
                        runs[0].font.color.rgb = _rgb(AI_SUGGESTION_COLOR["txt"])

    doc.add_paragraph("")

    if screenshot_path and os.path.exists(screenshot_path):
        try:
            doc.add_heading("📸 Ekran Görüntüsü", level=2)
            with PILImage.open(screenshot_path) as img:
                w_px, _ = img.size
            w_in = min(w_px / 96, 5.5)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(screenshot_path, width=Inches(w_in))
            cap = doc.add_paragraph(f"{page_name} sayfası ekran görüntüsü")
            cap.alignment         = WD_ALIGN_PARAGRAPH.CENTER
            cap.runs[0].font.size = Pt(9)
            cap.runs[0].italic    = True
        except Exception as e:
            print(f"⚠️  Ekran görüntüsü Word'e eklenemedi: {e}")

    doc.save(word_file)
    print(f"📄 Word kaydedildi: {word_file}")


# ── İç yardımcılar ───────────────────────────────────────────────────────────

def _build_ordered(all_elements: list, document_sections: list) -> list:
    """Elementleri config'deki section sırasına göre döndürür."""
    grouped = {s: [e for e in all_elements if e["status"] == s]
               for s in ALL_STATUSES}
    result = []
    for key in document_sections:
        result.extend(grouped.get(SECTION_TO_STATUS.get(key, ""), []))
    return result


def _get_val(elem: dict, key: str, elem_id: str, new_status: str) -> str:
    if key == "element_id":    return elem_id
    if key == "new_status":    return new_status
    if key == "ai_suggestion": return elem.get("ai_suggestion", "")
    return elem.get(key, "") or ""


def _style_excel_cell(c, key: str, pal: dict, r_fill, ns_fill, ai_fill, new_status: str) -> None:
    if key == "ai_suggestion":
        c.fill      = ai_fill
        c.font      = Font(size=8, color=AI_SUGGESTION_COLOR["txt"])
        c.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
    elif key == "new_status":
        c.fill = ns_fill
        if new_status:
            c.font = Font(bold=True, color=NEW_STATUS_COLOR["txt"], size=10)
        else:
            c.font = Font(size=10)
        c.alignment = CENTER
    elif key == "status":
        c.fill      = r_fill
        c.font      = Font(bold=True, color=pal["txt"], size=10)
        c.alignment = CENTER
    elif key == "element_id":
        c.fill      = r_fill
        c.font      = Font(bold=True, size=10)
        c.alignment = CENTER
    else:
        c.fill      = r_fill
        c.font      = Font(size=10)
        c.alignment = LEFT